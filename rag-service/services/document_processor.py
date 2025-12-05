import asyncio
import time
from pathlib import Path
from typing import List, Dict, Any
import logging
from docling_parse.pdf_parser import pdf_parser_v2
import json

from models.schemas import DocumentData, DocumentChunk

logger = logging.getLogger(__name__)

class DocumentProcessor:
    """Service for processing documents using Docling"""
    
    def __init__(self):
        self.supported_formats = ['.pdf', '.docx', '.txt', '.html', '.md', '.xlsx', '.xls', '.csv']
    
    async def process_document(
        self, 
        file_path: Path, 
        chunk_size: int = 1000, 
        chunk_overlap: int = 200
    ) -> Dict[str, Any]:
        """
        Process a document and extract text chunks using Docling
        
        Args:
            file_path: Path to the document file
            chunk_size: Maximum size of each text chunk
            chunk_overlap: Overlap between consecutive chunks
            
        Returns:
            Dictionary containing processed document data
        """
        start_time = time.time()
        file_extension = file_path.suffix.lower()
        
        if file_extension not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {file_extension}")
        
        logger.info(f"Processing document: {file_path.name}")
        
        try:
            # Extract text based on file type
            if file_extension == '.pdf':
                content, metadata = await self._process_pdf(file_path)
            elif file_extension == '.docx':
                content, metadata = await self._process_docx(file_path)
            elif file_extension == '.txt':
                content, metadata = await self._process_txt(file_path)
<<<<<<< HEAD
            elif file_extension == '.md':
                content, metadata = await self._process_markdown(file_path)
=======
>>>>>>> 743801bcc0d94f9953f34961b803df0b4769c53d
            elif file_extension in ['.xlsx', '.xls']:
                content, metadata = await self._process_excel(file_path)
            elif file_extension == '.csv':
                content, metadata = await self._process_csv(file_path)
            else:
                # For other formats, try basic text extraction
                content, metadata = await self._process_generic(file_path)
            
            # Create chunks with metadata
            chunks = self._create_chunks(content, chunk_size, chunk_overlap, metadata)
            
<<<<<<< HEAD
            # Write chunks to file for verification (debug mode)
            self._write_chunks_to_file(file_path, chunks, metadata)
            
=======
>>>>>>> 743801bcc0d94f9953f34961b803df0b4769c53d
            processing_time = time.time() - start_time
            
            return {
                "filename": file_path.name,
                "content": content,
                "chunks": chunks,
                "metadata": metadata,
                "processing_time": processing_time
            }
            
        except Exception as e:
            logger.error(f"Error processing document {file_path.name}: {e}")
            raise Exception(f"Failed to process document: {str(e)}")
    
    async def _process_pdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process PDF using Docling-parse with pypdf fallback"""
        # Try Docling first
        docling_success = False
        try:
            # Use the exact same approach as the working extractor
            parser = pdf_parser_v2(str(file_path))
            parser.load_document('doc1', str(file_path))
            result = parser.parse_pdf_from_key('doc1')
            
            # Extract text from all pages
            text_data = {
                "total_pages": len(result['pages']),
                "text_elements": [],
                "full_text": "",
                "pages": {}
            }
            
            pages_to_process = range(len(result['pages']))
            
            for page_idx in pages_to_process:
                if 0 <= page_idx < len(result['pages']):
                    page = result['pages'][page_idx]
                    sanitized_page = page['sanitized']
                    
                    page_text = ""
                    page_elements = []
                    
                    # Extract text from cells
                    if 'cells' in sanitized_page and 'data' in sanitized_page['cells']:
                        for cell_idx, cell_data in enumerate(sanitized_page['cells']['data']):
                            # Cell data is a list where text is at index 12
                            if len(cell_data) > 12 and cell_data[12]:
                                text_content = str(cell_data[12])
                                if text_content.strip():
                                    element = {
                                        "element_index": cell_idx,
                                        "text": text_content,
                                        "char_count": len(text_content),
                                        "word_count": len(text_content.split()),
                                        "page": page_idx,
                                        "element_type": "cell",
                                        "bbox": [cell_data[0], cell_data[1], cell_data[2], cell_data[3]] if len(cell_data) > 3 else None
                                    }
                                    page_elements.append(element)
                                    page_text += f"\n{text_content}"
                    
                    text_data["text_elements"].extend(page_elements)
                    text_data["full_text"] += page_text
                    
                    text_data["pages"][str(page_idx)] = {
                        "text": page_text,
                        "char_count": len(page_text),
                        "word_count": len(page_text.split()),
                        "elements": page_elements
                    }
            
            content = text_data["full_text"]
            
            metadata = {
                "file_type": "pdf",
                "pages_count": text_data["total_pages"],
                "total_elements": len(text_data["text_elements"]),
                "document_info": result.get('info', {}),
                "extraction_method": "docling-parse"
            }
            
            # Check if we actually got meaningful content
            if content.strip():
                docling_success = True
                logger.info(f"Extracted {len(text_data['text_elements'])} text elements from {text_data['total_pages']} pages using Docling")
            else:
                logger.warning(f"No text content extracted from PDF with Docling: {file_path.name}, falling back to pypdf")
            
        except Exception as docling_error:
            logger.warning(f"Docling processing failed for {file_path.name}: {docling_error}")
        
        # If Docling failed or extracted no content, try pypdf
        if not docling_success:
            try:
                logger.info("Falling back to pypdf extraction...")
                return await self._process_pdf_with_pypdf(file_path)
            except Exception as pypdf_error:
                logger.error(f"Both Docling and pypdf failed for PDF {file_path.name}: {pypdf_error}")
                raise Exception(f"Failed to process PDF with both Docling and pypdf: {str(pypdf_error)}")
        
        # Return Docling results if successful
        return content, metadata
    
    async def _process_pdf_with_pypdf(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process PDF using pypdf as fallback when Docling fails"""
        try:
            from pypdf import PdfReader
            
            reader = PdfReader(str(file_path))
            
            text_content = []
            total_elements = 0
            
            for page_num, page in enumerate(reader.pages, 1):
                try:
                    page_text = page.extract_text()
                    if page_text.strip():
                        text_content.append(f"=== Page {page_num} ===\n{page_text}")
                        total_elements += len(page_text.split())
                except Exception as e:
                    logger.warning(f"Error extracting text from page {page_num}: {e}")
                    continue
            
            content = "\n\n".join(text_content)
            
            if not content.strip():
                logger.warning(f"No text content extracted from PDF with pypdf: {file_path.name}")
                content = f"PDF document {file_path.name} - No readable text content found"
            
            metadata = {
                "file_type": "pdf",
                "pages_count": len(reader.pages),
                "total_elements": total_elements,
                "document_info": reader.metadata if reader.metadata else {},
                "extraction_method": "pypdf"
            }
            
            logger.info(f"Extracted text from {len(reader.pages)} pages using pypdf")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing PDF with pypdf {file_path.name}: {e}")
            raise Exception(f"Failed to process PDF with pypdf: {str(e)}")
    
    async def _process_docx(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process DOCX file"""
        try:
            from docx import Document as DocxDocument
            doc = DocxDocument(file_path)
            content_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    content_parts.append(paragraph.text.strip())
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        content_parts.append(" | ".join(row_text))
            
            content = "\n".join(content_parts)
            
            metadata = {
                "file_type": "docx",
                "paragraphs_count": len(doc.paragraphs),
                "tables_count": len(doc.tables),
                "extraction_method": "python-docx"
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing DOCX {file_path.name}: {e}")
            raise Exception(f"Failed to process DOCX: {str(e)}")
    
    async def _process_excel(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process Excel file (XLSX/XLS)"""
        try:
            from openpyxl import load_workbook
            from openpyxl.cell.cell import MergedCell
            
            workbook = load_workbook(filename=file_path, data_only=True)
            content_parts = []
            
            total_sheets = len(workbook.sheetnames)
            total_rows = 0
            total_cells = 0
            sheets_data = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_content = []
                sheet_rows = 0
                sheet_cells = 0
                
                # Add sheet header
                sheet_content.append(f"\n=== Sheet: {sheet_name} ===")
                
                # Get merged cells for reference and track which cells are merged (but not the top-left)
                merged_cell_coords = set()
                for merged_range in sheet.merged_cells.ranges:
                    # Get all cells in the merged range
                    for row in range(merged_range.min_row, merged_range.max_row + 1):
                        for col in range(merged_range.min_col, merged_range.max_col + 1):
                            cell_coord = sheet.cell(row, col).coordinate
                            # Don't skip the top-left cell
                            if cell_coord != sheet.cell(merged_range.min_row, merged_range.min_col).coordinate:
                                merged_cell_coords.add(cell_coord)
                
                # Process rows with data
                for row_idx, row in enumerate(sheet.iter_rows(min_row=1, values_only=False), 1):
                    row_data = []
                    row_empty = True
                    
                    for col_idx, cell in enumerate(row, 1):
                        # Skip if cell is part of a merged range (but not the top-left)
                        if isinstance(cell, MergedCell):
                            continue
                        
                        if cell.coordinate in merged_cell_coords:
                            continue
                        
                        # Get cell value
                        cell_value = None
                        if cell.data_type == 'f':  # Formula
                            # Try to get calculated value first
                            if cell.value is not None:
                                cell_value = str(cell.value)
                            # Also store the formula itself
                            formula_text = f"(Formula: {cell.formula})"
                            cell_value = cell_value if cell_value else formula_text
                        else:
                            cell_value = cell.value
                        
                        if cell_value is not None:
                            cell_str = str(cell_value).strip()
                            if cell_str:
                                row_data.append(cell_str)
                                row_empty = False
                    
                    if not row_empty:
                        # Join row data with tab separator for table structure
                        sheet_content.append("\t".join(row_data))
                        row_cell_count = len(row_data)
                        sheet_rows += 1
                        sheet_cells += row_cell_count
                        total_rows += 1
                        total_cells += row_cell_count
                
                sheets_data.append({
                    "name": sheet_name,
                    "rows": sheet_rows,
                    "cells": sheet_cells
                })
                
                content_parts.extend(sheet_content)
            
            content = "\n".join(content_parts)
            
            metadata = {
                "file_type": "excel",
                "workbook_format": "xlsx" if file_path.suffix == '.xlsx' else "xls",
                "total_sheets": total_sheets,
                "sheet_names": workbook.sheetnames,
                "total_rows": total_rows,
                "total_cells": total_cells,
                "sheets_data": sheets_data,
                "extraction_method": "openpyxl"
            }
            
            workbook.close()
            
            logger.info(f"Extracted {total_rows} rows from {total_sheets} sheets")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing Excel {file_path.name}: {e}")
            raise Exception(f"Failed to process Excel file: {str(e)}")
    
    async def _process_csv(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process CSV file"""
        try:
            import pandas as pd
            
            # Read CSV with automatic encoding detection
            try:
                df = pd.read_csv(file_path, encoding='utf-8')
            except UnicodeDecodeError:
                try:
                    df = pd.read_csv(file_path, encoding='latin-1')
                except UnicodeDecodeError:
                    df = pd.read_csv(file_path, encoding='cp1252', errors='ignore')
            
            content_parts = []
            
            # Get number of rows and columns
            total_rows = len(df)
            total_cols = len(df.columns)
            
            # Add header row
            headers = df.columns.tolist()
            content_parts.append("\t".join([str(h) for h in headers]))
            
            # Add data rows
            for index, row in df.iterrows():
                row_data = []
                for col in df.columns:
                    cell_value = row[col]
                    # Handle NaN values
                    if pd.isna(cell_value):
                        cell_value = ""
                    else:
                        cell_value = str(cell_value).strip()
                    row_data.append(cell_value)
                
                # Only add non-empty rows
                if any(cell for cell in row_data if cell):
                    content_parts.append("\t".join(row_data))
            
            content = "\n".join(content_parts)
            
            # Count non-empty cells
            non_empty_cells = df.notna().sum().sum()
            
            metadata = {
                "file_type": "csv",
                "total_rows": total_rows,
                "total_columns": total_cols,
                "column_names": headers,
                "non_empty_cells": int(non_empty_cells),
                "extraction_method": "pandas"
            }
            
            logger.info(f"Extracted {total_rows} rows and {total_cols} columns from CSV")
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing CSV {file_path.name}: {e}")
            raise Exception(f"Failed to process CSV file: {str(e)}")
    
    async def _process_txt(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            metadata = {
                "file_type": "txt",
                "char_count": len(content),
                "line_count": len(content.splitlines()),
                "extraction_method": "direct_read"
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing TXT {file_path.name}: {e}")
            raise Exception(f"Failed to process TXT: {str(e)}")
    
<<<<<<< HEAD
    async def _process_markdown(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process Markdown file with image extraction support"""
        try:
            from markdown_it import MarkdownIt
            import re
            import base64
            from io import BytesIO
            from PIL import Image
            
            # Read markdown content
            with open(file_path, 'r', encoding='utf-8') as f:
                content = f.read()
            
            # Parse markdown to extract structure
            md = MarkdownIt()
            tokens = md.parse(content)
            
            # Extract text content preserving structure
            text_parts = []
            images = []
            
            i = 0
            while i < len(tokens):
                token = tokens[i]
                
                if token.type == 'heading_open':
                    # Get heading level
                    level = token.tag[1] if len(token.tag) > 1 else '1'
                    # Heading text will be in next inline token
                    i += 1
                    if i < len(tokens) and tokens[i].type == 'inline':
                        if tokens[i].content:
                            # Add heading with markdown syntax
                            text_parts.append(f"\n{'#' * int(level)} {tokens[i].content}\n")
                    i += 1  # Skip heading_close
                elif token.type == 'paragraph_open':
                    i += 1  # Skip to inline
                    if i < len(tokens) and tokens[i].type == 'inline':
                        if tokens[i].content:
                            text_parts.append(f"{tokens[i].content}\n")
                    i += 1  # Skip paragraph_close
                elif token.type == 'inline' and token.content:
                    text_parts.append(token.content)
                    i += 1
                elif token.type == 'text':
                    # Skip standalone text tokens, handled by inline
                    i += 1
                elif token.type == 'bullet_list_open':
                    i += 1
                    # Process list items
                    while i < len(tokens) and tokens[i].type != 'bullet_list_close':
                        if tokens[i].type == 'list_item_open':
                            i += 1
                            if i < len(tokens) and tokens[i].type == 'inline':
                                if tokens[i].content:
                                    text_parts.append(f"- {tokens[i].content}\n")
                            i += 1
                        else:
                            i += 1
                    i += 1  # Skip bullet_list_close
                else:
                    i += 1
            
            # Build full text content
            full_text = ' '.join(text_parts)
            
            # If markdown parsing didn't yield good results, use original content
            if not full_text.strip() or len(full_text) < len(content) * 0.3:
                logger.info("Markdown parsing yielded limited content, using original text")
                full_text = content
            
            # Remove image data from text content to reduce chunk size
            # We'll replace image references with placeholders
            full_text_no_images = full_text
            
            # Extract images (base64 embedded images)
            base64_pattern = re.compile(r'!\[([^\]]*)\]\(data:image/(png|jpeg|jpg|gif|webp);base64,([A-Za-z0-9+/=]+)\)')
            
            for match in re.finditer(base64_pattern, content):
                alt_text = match.group(1)
                image_format = match.group(2)
                base64_data = match.group(3)
                
                # Replace image in text with a placeholder
                placeholder = f"[Image: {alt_text or 'embedded_image'}]"
                full_text_no_images = full_text_no_images.replace(match.group(0), placeholder)
                
                try:
                    # Decode to verify it's valid image data
                    image_bytes = base64.b64decode(base64_data)
                    image = Image.open(BytesIO(image_bytes))
                    
                    images.append({
                        "alt_text": alt_text,
                        "format": image_format,
                        "base64_data": base64_data,
                        "width": image.size[0],
                        "height": image.size[1],
                        "size_bytes": len(image_bytes),
                        "position_in_text": match.start()
                    })
                except Exception as e:
                    logger.warning(f"Failed to process embedded image in markdown: {e}")
            
            # Also check for HTML img tags
            html_img_pattern = re.compile(
                r'<img[^>]*src=["\'](data:image/(png|jpeg|jpg|gif|webp);base64,([A-Za-z0-9+/=]+))["\'][^>]*>',
                re.IGNORECASE
            )
            
            for match in re.finditer(html_img_pattern, content):
                full_data_uri = match.group(1)
                image_format = match.group(2)
                base64_data = match.group(3)
                
                # Extract alt text if present
                alt_match = re.search(r'alt=["\']([^"\']*)["\']', match.group(0), re.IGNORECASE)
                alt_text = alt_match.group(1) if alt_match else ""
                
                # Replace image in text with a placeholder
                placeholder = f"[Image: {alt_text or 'embedded_image'}]"
                full_text_no_images = full_text_no_images.replace(match.group(0), placeholder)
                
                try:
                    image_bytes = base64.b64decode(base64_data)
                    image = Image.open(BytesIO(image_bytes))
                    
                    images.append({
                        "alt_text": alt_text,
                        "format": image_format,
                        "base64_data": base64_data,
                        "width": image.size[0],
                        "height": image.size[1],
                        "size_bytes": len(image_bytes),
                        "position_in_text": match.start(),
                        "type": "html_img_tag"
                    })
                except Exception as e:
                    logger.warning(f"Failed to process HTML img tag in markdown: {e}")
            
            # Use text without images for chunking
            full_text = full_text_no_images
            
            metadata = {
                "file_type": "markdown",
                "char_count": len(content),
                "line_count": len(content.splitlines()),
                "image_count": len(images),
                "has_images": len(images) > 0,
                "images": images,
                "extraction_method": "markdown-it-py with image extraction"
            }
            
            logger.info(f"Extracted {len(images)} images from markdown file")
            
            return full_text, metadata
            
        except Exception as e:
            logger.error(f"Error processing markdown {file_path.name}: {e}")
            # Fallback to basic text read
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()
                
                metadata = {
                    "file_type": "markdown",
                    "char_count": len(content),
                    "line_count": len(content.splitlines()),
                    "extraction_method": "direct_read_fallback",
                    "has_images": False,
                    "image_count": 0
                }
                
                return content, metadata
            except Exception as fallback_error:
                logger.error(f"Fallback processing also failed: {fallback_error}")
                raise Exception(f"Failed to process markdown: {str(e)}")
    
=======
>>>>>>> 743801bcc0d94f9953f34961b803df0b4769c53d
    async def _process_generic(self, file_path: Path) -> tuple[str, Dict[str, Any]]:
        """Process other file types with basic text extraction"""
        try:
            # Try to read as text file
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read()
            
            metadata = {
                "file_type": file_path.suffix.lower(),
                "char_count": len(content),
                "extraction_method": "generic_text_read"
            }
            
            return content, metadata
            
        except Exception as e:
            logger.error(f"Error processing generic file {file_path.name}: {e}")
            raise Exception(f"Failed to process file: {str(e)}")
    
    def _create_chunks(self, text: str, chunk_size: int, chunk_overlap: int, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
<<<<<<< HEAD
        """
        Split text into overlapping chunks using LangChain's RecursiveCharacterTextSplitter
        
        Args:
            text: Text to chunk
            chunk_size: Size of each chunk (in characters)
            chunk_overlap: Overlap between chunks (in characters)
            metadata: Additional metadata to include
            
        Returns:
            List of chunk dictionaries with text and metadata
        """
        if not text.strip():
            return []
        
        try:
            from langchain_text_splitters import RecursiveCharacterTextSplitter
            
            # Create LangChain text splitter with custom parameters
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=chunk_size,
                chunk_overlap=chunk_overlap,
                length_function=len,
                separators=[
                    "\n\n",  # Paragraphs
                    "\n",    # Lines
                    ". ",    # Sentences
                    "! ",    # Exclamations
                    "? ",    # Questions
                    ", ",    # Commas
                    " ",     # Words
                    ""       # Characters (fallback)
                ]
            )
            
            # Split text using LangChain
            langchain_chunks = text_splitter.create_documents([text])
            
            # Convert to our format
            chunks = []
            for idx, chunk in enumerate(langchain_chunks):
                chunk_text = chunk.page_content
                
                # Get character position info
                start_char = chunk.metadata.get('start_index', idx * chunk_size)
                end_char = chunk.metadata.get('end_index', start_char + len(chunk_text))
                
                # Estimate page number based on position
                total_chars = len(text)
                estimated_pages = max(1, total_chars // 2000)
                estimated_page = min(estimated_pages, max(1, int((start_char / total_chars) * estimated_pages) + 1))
                
                chunks.append({
                    "text": chunk_text,
                    "chunk_index": idx,
                    "start_char": start_char,
                    "end_char": end_char,
                    "metadata": {
                        "chunk_size": len(chunk_text),
                        "char_start": start_char,
                        "char_end": end_char,
                        "page_number": estimated_page,
                        **(metadata or {})
                    }
                })
            
            logger.info(f"Created {len(chunks)} chunks using LangChain RecursiveCharacterTextSplitter")
            return chunks
            
        except ImportError:
            logger.warning("LangChain not available, falling back to basic chunking")
            return self._create_chunks_basic(text, chunk_size, chunk_overlap, metadata)
        
        except Exception as e:
            logger.error(f"Error creating chunks with LangChain: {e}, falling back to basic chunking")
            return self._create_chunks_basic(text, chunk_size, chunk_overlap, metadata)
    
    def _create_chunks_basic(self, text: str, chunk_size: int, chunk_overlap: int, metadata: Dict[str, Any] = None) -> List[Dict[str, Any]]:
        """Fallback basic chunking method when LangChain is not available"""
=======
        """Split text into overlapping chunks with page information"""
>>>>>>> 743801bcc0d94f9953f34961b803df0b4769c53d
        if not text.strip():
            return []

        chunks = []
        start = 0
        chunk_index = 0
        
        # Estimate page numbers based on text length (rough approximation)
        total_chars = len(text)
        estimated_pages = max(1, total_chars // 2000)  # Assume ~2000 chars per page
        
        while start < len(text):
            end = start + chunk_size
            
            # Try to break at sentence boundary
            if end < len(text):
                # Look for sentence endings within the last 100 characters
                search_start = max(start, end - 100)
                sentence_end = text.rfind('.', search_start, end)
                if sentence_end > start:
                    end = sentence_end + 1
                else:
                    # Look for word boundary
                    word_end = text.rfind(' ', search_start, end)
                    if word_end > start:
                        end = word_end
            
            chunk_text = text[start:end].strip()
            
            if chunk_text:
                # Estimate page number for this chunk
                estimated_page = min(estimated_pages, max(1, (start / total_chars) * estimated_pages + 1))
                
                chunks.append({
                    "text": chunk_text,
                    "chunk_index": chunk_index,
                    "start_char": start,
                    "end_char": end,
                    "metadata": {
                        "chunk_size": len(chunk_text),
                        "char_start": start,
                        "char_end": end,
                        "page_number": int(estimated_page),
                        **(metadata or {})
                    }
                })
                chunk_index += 1
            
            # Move start position with overlap
            start = end - chunk_overlap
            if start >= len(text):
                break
        
<<<<<<< HEAD
        logger.info(f"Created {len(chunks)} chunks using basic chunking")
        return chunks
    
    def _write_chunks_to_file(self, file_path: Path, chunks: List[Dict[str, Any]], metadata: Dict[str, Any]):
        """Write chunks to a file for verification"""
        try:
            # Create chunks directory if it doesn't exist
            chunks_dir = Path("chunks_debug")
            chunks_dir.mkdir(exist_ok=True)
            
            # Create output filename
            base_name = file_path.stem
            output_file = chunks_dir / f"{base_name}_chunks.txt"
            
            with open(output_file, 'w', encoding='utf-8') as f:
                f.write(f"Document: {file_path.name}\n")
                f.write(f"File Type: {metadata.get('file_type', 'unknown')}\n")
                f.write(f"Total Chunks: {len(chunks)}\n")
                if 'image_count' in metadata:
                    f.write(f"Images Found: {metadata.get('image_count', 0)}\n")
                f.write("=" * 80 + "\n\n")
                
                for i, chunk in enumerate(chunks):
                    f.write(f"--- Chunk {i} ---\n")
                    f.write(f"Chunk Index: {chunk.get('chunk_index', i)}\n")
                    f.write(f"Start Char: {chunk.get('start_char', 'N/A')}\n")
                    f.write(f"End Char: {chunk.get('end_char', 'N/A')}\n")
                    if 'metadata' in chunk and 'page_number' in chunk['metadata']:
                        f.write(f"Page: {chunk['metadata']['page_number']}\n")
                    f.write(f"Text Length: {len(chunk.get('text', ''))}\n")
                    f.write("\nText Content:\n")
                    f.write(chunk.get('text', '')[:500])  # First 500 chars
                    if len(chunk.get('text', '')) > 500:
                        f.write(f"\n... (truncated, total {len(chunk.get('text', ''))} chars)")
                    f.write("\n\n")
                
                if 'images' in metadata and metadata['images']:
                    f.write("\n" + "=" * 80 + "\n")
                    f.write("EXTRACTED IMAGES:\n")
                    f.write("=" * 80 + "\n")
                    for i, img in enumerate(metadata['images']):
                        f.write(f"\nImage {i+1}:\n")
                        f.write(f"  Alt Text: {img.get('alt_text', 'N/A')}\n")
                        f.write(f"  Format: {img.get('format', 'N/A')}\n")
                        f.write(f"  Dimensions: {img.get('width', 'N/A')}x{img.get('height', 'N/A')}\n")
                        f.write(f"  Size: {img.get('size_bytes', 0):,} bytes\n")
                        if 'position_in_text' in img:
                            f.write(f"  Position in text: {img['position_in_text']}\n")
            
            logger.info(f"Chunks written to: {output_file}")
            
        except Exception as e:
            logger.warning(f"Failed to write chunks to file: {e}")
            # Don't fail the whole process if debug logging fails
=======
        logger.info(f"Created {len(chunks)} chunks from document")
        return chunks
>>>>>>> 743801bcc0d94f9953f34961b803df0b4769c53d
